import re
import time
import traceback
import os
from datetime import datetime
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import NoSuchElementException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def get_chrome_driver():
    chrome_service = Service(ChromeDriverManager().install())
    chrome_options = webdriver.ChromeOptions()
    return webdriver.Chrome(service=chrome_service, options=chrome_options)

def scroll_down(driver, times):
    for _ in range(times):
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)

def extract_price(price_text):
    """Extract numerical value from price text"""
    try:
        return float(re.sub(r'[^\d.]', '', price_text).replace(',', ''))
    except ValueError:
        return None

def extract_discount_rate(rate_text):
    """Extract numerical value from discount rate text"""
    try:
        return float(re.sub(r'[^\d.]', '', rate_text))
    except ValueError:
        return None

def get_product_info(driver, category, user_brand_name):
    products = []
    unique_prices = set()

    try:
        product_elements = driver.find_elements(By.XPATH, "//div[contains(@class, 'sc-widb61-1 leHifG')]")

        if not product_elements:
            products.append({
                "Category": category,
                "Brand Name": user_brand_name,
                "Product Name": "X",
                "Original Price": None,
                "Discount Rate": None,
                "Actual Sale Price": None
            })
            print(f"No products found for brand {user_brand_name} and category {category}.")
            return products

        for element in product_elements:
            # 상품명
            try:
                product_name = element.find_element(By.XPATH, ".//span[contains(@class, 'sc-dIMoHT')]").text
            except NoSuchElementException:
                product_name = "N/A"

            # 실제 판매 가격
            try:
                sale_price_text = element.find_element(By.XPATH, ".//span[contains(@class, 'sc-fWnslK') and not(contains(@class, 'text-red'))]").text
                actual_sale_price = extract_price(sale_price_text)
            except NoSuchElementException:
                actual_sale_price = None

            if actual_sale_price in unique_prices:
                continue
            unique_prices.add(actual_sale_price)

            # 할인율 존재 여부 체크
            try:
                discount_rate_text = element.find_element(By.XPATH, ".//span[contains(@class, 'sc-fWnslK') and contains(@class, 'text-red')]").text
                discount_rate = extract_discount_rate(discount_rate_text)
            except NoSuchElementException:
                discount_rate = None

            # 원래 가격(정가) 설정
            original_price = None
            if discount_rate is not None:
                try:
                    # 정가와 할인가가 같이 표시될 경우, 실제 판매가 다음에 있는 가격이 정가로 간주됩니다.
                    price_elements = element.find_elements(By.XPATH, ".//span[contains(@class, 'sc-fWnslK')]")
                    for price_el in price_elements:
                        price = price_el.text
                        if price and price != sale_price_text and '원' in price:
                            original_price = extract_price(price)
                            break
                except NoSuchElementException:
                    original_price = actual_sale_price
            else:
                original_price = actual_sale_price

            products.append({
                "Category": category,
                "Brand Name": user_brand_name,
                "Product Name": product_name,
                "Original Price": original_price,
                "Discount Rate": discount_rate,
                "Actual Sale Price": actual_sale_price
            })

    except Exception as e:
        print(f"Error occurred while fetching product info: {e}")
        print(traceback.format_exc())

    return products

def calculate_summary_stats(df):
    summary = []
    category_stats = []

    for (category, brand_name), group in df.groupby(['Category', 'Brand Name']):
        price_ori_stats = group['Original Price'].dropna()
        actual_price_stats = group['Actual Sale Price'].dropna()
        discount_stats = group['Discount Rate'].dropna()

        if not price_ori_stats.empty:
            max_ori = price_ori_stats.max()
            min_ori = price_ori_stats.min()
            avg_ori = price_ori_stats.mean()
            avg_ori_max_min = (max_ori + min_ori) / 2
            summary.append({
                "Category": category,
                "Brand Name": brand_name,
                "Metric": f"Original Price ({len(price_ori_stats)})",
                "Max": max_ori,
                "Min": min_ori,
                "Avg Max Min": avg_ori_max_min,
                "Avg": avg_ori
            })

        if not actual_price_stats.empty:
            max_actual = actual_price_stats.max()
            min_actual = actual_price_stats.min()
            avg_actual = actual_price_stats.mean()
            avg_actual_max_min = (max_actual + min_actual) / 2
            avg_discount_rate = discount_stats.mean() if not discount_stats.empty else 0

            summary.append({
                "Category": category,
                "Brand Name": brand_name,
                "Metric": f"Actual Sale Price ({avg_discount_rate:.2f}%)",
                "Max": max_actual,
                "Min": min_actual,
                "Avg Max Min": avg_actual_max_min,
                "Avg": avg_actual
            })

    for category, group in df.groupby('Category'):
        ori_prices = group['Original Price'].dropna()
        actual_prices = group['Actual Sale Price'].dropna()

        if not ori_prices.empty:
            brand_name = f"ALL ({category})"
            category_stats.append({
                "Category": category,
                "Brand Name": brand_name,
                "Metric": f"Original Price ({len(ori_prices)})",
                "Max": ori_prices.max(),
                "Min": ori_prices.min(),
                "Avg Max Min": (ori_prices.max() + ori_prices.min()) / 2,
                "Avg": ori_prices.mean()
            })

        if not actual_prices.empty:
            avg_discount_rate = group['Discount Rate'].mean() if not group['Discount Rate'].empty else 0
            category_stats.append({
                "Category": category,
                "Brand Name": brand_name,
                "Metric": f"Actual Sale Price ({avg_discount_rate:.2f}%)",
                "Max": actual_prices.max(),
                "Min": actual_prices.min(),
                "Avg Max Min": (actual_prices.max() + actual_prices.min()) / 2,
                "Avg": actual_prices.mean()
            })

    summary_df = pd.DataFrame(summary)
    category_summary_df = pd.DataFrame(category_stats).sort_values(by='Category')

    return summary_df, category_summary_df

def generate_unique_filename(base_filename):
    if not os.path.exists(base_filename):
        return base_filename
    base, ext = os.path.splitext(base_filename)
    timestamp = datetime.now().strftime("%H%M%S")
    return f"{base}_{timestamp}{ext}"

def format_columns_as_numbers(df, columns):
    for column in columns:
        df[column] = df[column].apply(lambda x: float(re.sub(r'[^\d.]', '', str(x))) if pd.notnull(x) else x)
        df[column] = df[column].astype(int)
    return df

def plot_bar_chart(data, category, filename):
    """Generate bar chart for each category and save to file."""
    plt.figure(figsize=(10, 5))
    brands = data['Brand Name'].unique()
    index = range(len(brands))

    original_prices = [
        data[(data['Brand Name'] == brand) & (data['Metric'].str.contains('Original Price'))]['Avg'].values[0]
        if not data[(data['Brand Name'] == brand) & (data['Metric'].str.contains('Original Price'))].empty else 0
        for brand in brands
    ]
    actual_prices = [
        data[(data['Brand Name'] == brand) & (data['Metric'].str.contains('Actual Sale Price'))]['Avg'].values[0]
        if not data[(data['Brand Name'] == brand) & (data['Metric'].str.contains('Actual Sale Price'))].empty else 0
        for brand in brands
    ]

    plt.bar(index, original_prices, width=0.4, label='Original Price', align='center')
    plt.bar(index, actual_prices, width=0.4, label='Actual Sale Price', align='edge')

    plt.xlabel('Brand')
    plt.ylabel('Average Price')
    plt.title(f'Average Prices for {category}')
    plt.xticks(index, brands, rotation=45, ha="right")
    plt.legend()
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

def create_docx_report(summary_df, report_filename):
    doc = Document()
    doc.add_heading('Product Price Analysis Report', 0)

    for category in summary_df['Category'].unique():
        doc.add_heading(f'Category: {category}', level=1)
        category_data = summary_df[summary_df['Category'] == category]

        brands = category_data['Brand Name'].unique()
        for brand in brands:
            brand_data = category_data[category_data['Brand Name'] == brand]
            doc.add_heading(f'Brand: {brand}', level=2)

            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric (Count or Rate)'
            hdr_cells[1].text = 'Max'
            hdr_cells[2].text = 'Min'
            hdr_cells[3].text = 'Avg Max Min'
            hdr_cells[4].text = 'Avg'

            for _, row in brand_data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row['Metric']
                row_cells[1].text = str(row['Max'])
                row_cells[2].text = str(row['Min'])
                row_cells[3].text = str(row['Avg Max Min'])
                row_cells[4].text = str(row['Avg'])

        # Add bar chart
        chart_filename = f"chart_{category}.png"
        plot_bar_chart(category_data, category, chart_filename)
        doc.add_picture(chart_filename, width=Inches(6.0))

    doc.save(report_filename)

def main():
    titles = input("Enter the product titles separated by commas: ").split(',')
    brand_names = input("Enter the brand names separated by commas: ").split(',')
    driver = get_chrome_driver()
    all_products = []

    try:
        for brand_name in brand_names:
            user_brand_name = brand_name.strip()
            for title in titles:
                title = title.strip()
                url = f"https://www.musinsa.com/brands/{user_brand_name}?includeKeywords={title}&sortCode=NEW&page=1&size=120&listViewType=3GridView"
                driver.get(url)
                time.sleep(3)
                scroll_down(driver, 2)
                products = get_product_info(driver, title, user_brand_name)
                all_products.extend(products)
    except WebDriverException as e:
        print(f"WebDriverException encountered: {e}")
    finally:
        driver.quit()

    if not all_products:
        print("No products found.")
        return

    df = pd.DataFrame(all_products)
    df.sort_values(by=['Category', 'Brand Name'], inplace=True)

    summary_df, category_summary_df = calculate_summary_stats(df)
    summary_df = format_columns_as_numbers(summary_df, ['Max', 'Min', 'Avg Max Min', 'Avg'])
    category_summary_df = format_columns_as_numbers(category_summary_df, ['Max', 'Min', 'Avg Max Min', 'Avg'])

    today_date = datetime.now().strftime("%Y%m%d")
    base_filename = f'musinsa_products_{today_date}.xlsx'
    filename = generate_unique_filename(base_filename)

    with pd.ExcelWriter(filename, engine='xlsxwriter') as writer:
        workbook = writer.book
        light_format = workbook.add_format({'bg_color': '#EEEEEE', 'border': 1})

        category_summary_df.to_excel(writer, index=False, sheet_name='Summary')
        category_summary_worksheet = writer.sheets['Summary']

        for i, col in enumerate(category_summary_df.columns):
            max_len = category_summary_df[col].astype(str).map(len).max()
            category_summary_worksheet.set_column(i, i, max_len + 2)
        category_summary_worksheet.autofilter(0, 0, len(category_summary_df), len(category_summary_df.columns) - 1)

        start_row = len(category_summary_df) + 2
        summary_df.to_excel(writer, index=False, sheet_name='Summary', startrow=start_row)
        summary_worksheet = writer.sheets['Summary']

        for i, col in enumerate(summary_df.columns):
            max_len = summary_df[col].astype(str).map(len).max()
            summary_worksheet.set_column(i, i, max_len + 2)
        summary_worksheet.autofilter(start_row, 0, start_row + len(summary_df), len(summary_df.columns) - 1)

        categories = summary_df['Category'].unique()
        color_scale = ['#F0F0F0', '#F8F8F8']

        for idx, category in enumerate(categories):
            category_rows = summary_df[summary_df['Category'] == category].index + start_row
            format = workbook.add_format({'bg_color': color_scale[idx % len(color_scale)], 'border': 1})

            for row in category_rows:
                summary_worksheet.set_row(row, cell_format=format)

        df.to_excel(writer, index=False, sheet_name='Products')
        products_worksheet = writer.sheets['Products']
        for i, col in enumerate(df.columns):
            max_len = df[col].astype(str).map(len).max()
            products_worksheet.set_column(i, i, max_len + 2)
        products_worksheet.autofilter(0, 0, len(df), len(df.columns) - 1)

    print(f"Data saved to {filename}")

    report_filename = f'report_{today_date}.docx'
    create_docx_report(summary_df, report_filename)
    print(f"Report saved to {report_filename}")

if __name__ == "__main__":
    main()
